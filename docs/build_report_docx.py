from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "BaoCao_HOG_SVM.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        element = tc_mar.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("BÁO CÁO DỰ ÁN CUỐI KỲ")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Phát Hiện Người Đi Bộ Bằng Image Gradient, HOG và SVM")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    run = p.add_run("Chủ đề 2: Image Gradient & HOG + SVM - Phát hiện đối tượng")
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    rows = [
        ("Môn học", "Nhập môn Xử lý ảnh số"),
        ("Ngôn ngữ / thư viện", "Python, OpenCV, scikit-image, scikit-learn, joblib"),
        ("Demo chính", "Realtime pedestrian detection trên video/webcam (OpenCV Pretrained)"),
        ("Model mở rộng", "Custom Linear SVM train từ 300 positive và 300 negative INRIA subset"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2300, 7060])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_GRAY)
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)

    doc.add_page_break()


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Báo cáo cuối kỳ Xử lý ảnh số: HOG + SVM Pedestrian Detection"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Trang ")
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_dataset_table(doc):
    data = [
        ("Nguồn train", "INRIA Person mirror trên Hugging Face", "Tập dữ liệu ngoài chuẩn, tách biệt hoàn toàn khỏi video demo"),
        ("Video đầu vào", "dataset/videos/walking.mp4", "Chỉ sử dụng làm video nguồn để chạy demo"),
        ("Positive (Dương tính)", "300 ảnh chuẩn hóa 64x128", "Chứa người đi bộ, tải từ data_ped/pedestrians"),
        ("Negative (Âm tính)", "300 ảnh chuẩn hóa 64x128", "Ảnh nền không có người, được cắt ngẫu nhiên từ data_ped/no_pedestrians"),
        ("Test images", "5 ảnh", "Dùng để kiểm thử nhanh và xuất ảnh Image Gradient, HOG"),
        ("Custom model", "models/custom_hog_svm.pkl", "Mô hình Linear SVM lưu trữ sau khi huấn luyện thành công"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Thành phần", "Số lượng / Đường dẫn", "Vai trò thực tế"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, BLUE_GRAY)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    set_repeat_table_header(table.rows[0])
    for row in data:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    set_table_geometry(table, [1900, 3300, 4160])


def add_result_table(doc):
    data = [
        ("Train/Test split", "480 mẫu train (80%), 120 mẫu test (20%)", "Phân chia ngẫu nhiên phân tầng (Stratified Split)"),
        ("Độ chính xác Test (Accuracy)", "96.67%", "Đánh giá khả năng tổng quát hóa trên tập kiểm thử độc lập"),
        ("Ma trận nhầm lẫn Test", "[[59, 1], [3, 57]]", "Chỉ có 1 mẫu nền bị nhầm là người, 3 mẫu người bị bỏ sót"),
        ("Độ chính xác toàn bộ", "99.33%", "Đánh giá lại trên toàn bộ 600 mẫu dữ liệu thu thập"),
        ("Ma trận nhầm lẫn toàn bộ", "[[299, 1], [3, 297]]", "Chứng minh mô hình tự huấn luyện học rất tốt các đặc trưng HOG"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Chỉ số", "Kết quả thực nghiệm", "Ý nghĩa / Diễn giải chi tiết"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, BLUE_GRAY)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    set_repeat_table_header(table.rows[0])
    for row in data:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    set_table_geometry(table, [2300, 2600, 4460])


def add_comparison_table(doc):
    data = [
        ("Nguồn gốc SVM", "Huấn luyện trên tập INRIA gốc quy mô lớn của OpenCV", "Huấn luyện trên subset 600 ảnh do nhóm tự tải về từ Hugging Face"),
        ("Kỹ thuật quét ảnh", "Sử dụng thư viện tối ưu hóa bằng C++ của OpenCV (detectMultiScale)", "Quét cửa sổ trượt (Sliding Window) đa tỷ lệ lập trình bằng Python"),
        ("Tốc độ khung hình (FPS)", "Cao (15 - 28 FPS), chạy thời gian thực rất mượt mà", "Thấp (~0.5 - 2 FPS), bị gián đoạn do vòng lặp quét thuần Python"),
        ("Độ chính xác thực tế", "Nhận diện tốt người đi bộ ở nhiều khoảng cách, ít bỏ sót", "Nhận diện tốt ở cự ly gần và trung bình, nhạy cảm với tỷ lệ quét"),
        ("Mục đích sử dụng", "Dùng để chạy demo chính thời gian thực trên webcam/video", "Chứng minh quy trình từ xử lý ảnh thô đến học máy truyền thống"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Đặc điểm so sánh", "OpenCV Pretrained Detector", "Custom Sliding-Window Detector"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, BLUE_GRAY)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    set_repeat_table_header(table.rows[0])
    for row in data:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    set_table_geometry(table, [2300, 3530, 3530])


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(9)


def add_figure(doc, image_path, caption, width=5.7):
    path = ROOT / image_path
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def build_report():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    add_title_page(doc)

    # CHƯƠNG 1
    doc.add_heading("CHƯƠNG 1: GIỚI THIỆU", level=1)
    
    doc.add_heading("1.1 Đặt vấn đề và tính cấp thiết của đề tài", level=2)
    doc.add_paragraph(
        "Phát hiện người đi bộ (Pedestrian Detection) là một trong những bài toán kinh điển và quan trọng nhất "
        "trong lĩnh vực Thị giác máy tính (Computer Vision) và Xử lý ảnh (Image Processing). Bài toán này đóng vai trò "
        "nền tảng trong nhiều ứng dụng thực tiễn quan trọng, bao gồm:"
    )
    add_bullets(doc, [
        "Hệ thống giám sát an ninh thông minh: Tự động phát hiện và theo dõi sự xâm nhập trái phép của con người vào các khu vực cấm hoặc giám sát mật độ đám đông ở nơi công cộng.",
        "Hệ thống hỗ trợ lái xe tiên tiến (ADAS) và Xe tự hành: Tự động nhận diện người đi bộ trên đường để đưa ra các cảnh báo phanh khẩn cấp hoặc điều khiển xe tránh va chạm, trực tiếp bảo vệ tính mạng con người.",
        "Phân tích hành vi khách hàng: Trong các trung tâm thương mại hoặc cửa hàng bán lẻ, việc phát hiện và đếm số lượng người giúp đánh giá mật độ di chuyển và thói quen mua sắm."
    ])
    doc.add_paragraph(
        "Mặc dù các phương pháp Học sâu (Deep Learning) như YOLO, SSD hay Faster R-CNN đang chiếm ưu thế nhờ độ chính xác vượt trội, "
        "các phương pháp thị giác máy tính truyền thống (Traditional Computer Vision) kết hợp giữa trích xuất đặc trưng thủ công "
        "(Hand-crafted features) và học máy cổ điển (Classical Machine Learning) vẫn giữ nguyên giá trị khoa học to lớn. Trong đó, "
        "sự kết hợp giữa Histogram of Oriented Gradients (HOG) và Support Vector Machine (SVM) (được giới thiệu lần đầu bởi Navneet "
        "Dalal và Bill Triggs vào năm 2005) là một cột mốc lịch sử quan trọng. Kỹ thuật này nổi bật nhờ tính toán hiệu quả, cấu trúc "
        "toán học chặt chẽ, dễ giải thích và không yêu cầu năng lực tính toán cực lớn (như GPU) để huấn luyện hoặc chạy thử nghiệm."
    )

    doc.add_heading("1.2 Mục tiêu đề tài", level=2)
    doc.add_paragraph("Dự án được xây dựng nhằm đạt được các mục tiêu sau:")
    add_numbered(doc, [
        "Nghiên cứu lý thuyết: Am hiểu sâu sắc bản chất toán học của toán tử vi phân tính toán Image Gradient, nguyên lý hoạt động của bộ mô tả đặc trưng HOG và cơ sở tối ưu hóa phân lớp của SVM.",
        "Xây dựng chương trình Demo: Phát triển một phần mềm hoàn chỉnh bằng Python và OpenCV có khả năng đọc luồng video từ file hoặc trực tiếp từ Webcam, thực hiện phát hiện người đi bộ theo thời gian thực và vẽ bounding box, hiển thị chỉ số FPS cũng như độ tin cậy.",
        "Mở rộng & Sáng tạo: Tự thiết kế quy trình tải dữ liệu ngoài từ Hugging Face, tự động tiền xử lý và cắt trích các mẫu ảnh, huấn luyện một mô hình Custom Linear SVM để phát hiện người đi bộ, thay vì chỉ sử dụng mô hình có sẵn của OpenCV.",
        "Trực quan hóa: Xuất các kết quả xử lý trung gian bao gồm ảnh Gradient Magnitude, Gradient Orientation và bản đồ HOG để phục vụ việc giải thích cơ chế vật lý của các thuật toán xử lý ảnh số."
    ])

    # CHƯƠNG 2
    doc.add_heading("CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ GIẢI THUẬT", level=1)
    
    doc.add_heading("2.1 Toán tử vi phân và Image Gradient (Độ dốc ảnh)", level=2)
    doc.add_paragraph(
        "Gradient của một ảnh biểu thị sự thay đổi cường độ sáng theo các hướng trong không gian 2 chiều. "
        "Đối với một ảnh xám liên tục I(x, y), gradient tại tọa độ (x, y) là một vector hướng từ vùng tối sang "
        "vùng sáng và được định nghĩa bởi các đạo hàm riêng:\n"
        "∇I(x, y) = [Gx, Gy]^T = [∂I/∂x, ∂I/∂y]^T"
    )
    doc.add_paragraph(
        "Trong xử lý ảnh số (miền rời rạc), các đạo hàm riêng này được xấp xỉ bằng phép toán tích chập (convolution) "
        "ảnh với các mặt nạ vi phân (differential kernels). Trong dự án này, nhóm sử dụng Toán tử Sobel kích thước 3x3 "
        "để tính toán các thành phần gradient Gx và Gy:"
    )
    doc.add_paragraph("- Mặt nạ Sobel theo trục X (Sx) để phát hiện biên dọc:")
    add_code_block(doc, "Sx = [[-1,  0,  1],\n      [-2,  0,  2],\n      [-1,  0,  1]]\nGx(x, y) = I(x, y) * Sx")
    doc.add_paragraph("- Mặt nạ Sobel theo trục Y (Sy) để phát hiện biên ngang:")
    add_code_block(doc, "Sy = [[-1, -2, -1],\n      [ 0,  0,  0],\n      [ 1,  2,  1]]\nGy(x, y) = I(x, y) * Sy")
    doc.add_paragraph("Từ hai thành phần này, ta tính được biên độ (Magnitude) và hướng (Orientation) của gradient tại mỗi điểm ảnh:")
    add_code_block(doc, "Biên độ Gradient:\nM(x, y) = sqrt(Gx(x, y)^2 + Gy(x, y)^2)\n\nHướng Gradient:\ntheta(x, y) = atan2(Gy(x, y), Gx(x, y))")
    doc.add_paragraph(
        "Hướng theta được ánh xạ về khoảng [0, 180 độ] (unsigned gradient) trong bài toán trích xuất HOG "
        "để tránh ảnh hưởng của việc đổi màu sáng - tối ngược hướng."
    )

    doc.add_heading("2.2 Đặc trưng Histogram of Oriented Gradients (HOG)", level=2)
    doc.add_paragraph(
        "Mục đích của HOG là mô tả hình dáng và cấu trúc bề ngoài của đối tượng bằng sự phân bố của hướng gradient "
        "trong các vùng ảnh cục bộ. Quy trình trích xuất HOG bao gồm các bước chặt chẽ sau:"
    )
    add_numbered(doc, [
        "Tiền xử lý và Chuẩn hóa độ sáng ảnh: Cửa sổ dò tìm chuẩn có kích thước 64x128 pixel. Áp dụng phép biến đổi căn bậc hai (transform_sqrt=True) trên ảnh để giảm bớt độ nhạy cảm đối với các biến đổi ánh sáng môi trường hay bóng đổ.",
        "Tính toán Gradient: Sử dụng các toán tử lọc 1D đơn giản [-1, 0, 1] và [-1, 0, 1]^T để tính Gx và Gy. Thực nghiệm cho thấy các bộ lọc vi phân nhỏ này giữ lại thông tin chi tiết biên tốt hơn.",
        "Phân chia ô (Cells) và Tạo Histogram: Chia ảnh 64x128 thành các cell có kích thước 8x8 pixel (gồm 8x16 = 128 cells). Tại mỗi cell, ta thống kê lược đồ hướng gradient gồm 9 bins tương ứng từ 0 đến 180 độ. Giá trị vote được trọng số hóa bởi biên độ gradient M(x, y).",
        "Gom nhóm khối (Blocks) và Chuẩn hóa độ sáng cục bộ (Block Normalization): Các cell lân cận được gom lại thành các block có kích thước 2x2 cells (16x16 pixel), trượt chồng chập lên nhau với bước trượt là 1 cell (8 pixel). Tổng cộng có 7x15 = 105 blocks. Mỗi block có 2x2x9 = 36 phần tử. Áp dụng chuẩn hóa L2-Hys (L2-norm với Hysteresis) để khử nhiễu ánh sáng cục bộ.",
        "Tạo Vector đặc trưng cuối cùng: Ghép toàn bộ vector của 105 blocks lại với nhau tạo thành vector đặc trưng HOG duy nhất có kích thước: 105 blocks * 36 giá trị/block = 3780 chiều."
    ])

    doc.add_heading("2.3 Bộ phân loại Support Vector Machine (SVM)", level=2)
    doc.add_paragraph(
        "Support Vector Machine là thuật toán học máy có giám sát tìm kiếm một siêu phẳng (hyperplane) phân tách tối ưu "
        "giữa hai lớp dữ liệu sao cho khoảng cách lề (margin) giữa các điểm dữ liệu gần nhất của hai lớp (các vector hỗ trợ) "
        "đến siêu phẳng là lớn nhất. Trong bài toán phát hiện người đi bộ sử dụng đặc trưng HOG, siêu phẳng phân hoạch là siêu phẳng "
        "tuyến tính (Linear SVM) nhằm đảm bảo tốc độ tính toán nhanh. Hàm quyết định phân lớp có dạng:\n"
        "f(x) = w^T x + b"
    )
    doc.add_paragraph(
        "Trong đó x là vector đặc trưng HOG (3780 chiều), w là vector trọng số tối ưu hóa cần học, và b là bias. "
        "Nếu f(x) >= 0, mẫu ảnh được phân lớp là người đi bộ (+1, Positive); ngược lại là nền (0, Negative). "
        "Khoảng cách từ điểm dữ liệu đến siêu phẳng quyết định độ tin cậy (confidence score) của phép dự đoán."
    )

    # CHƯƠNG 3
    doc.add_heading("CHƯƠNG 3: THIẾT KẾ HỆ THỐNG VÀ ỨNG DỤNG DEMO", level=1)
    
    doc.add_heading("3.1 Cấu trúc mã nguồn của dự án", level=2)
    doc.add_paragraph("Mã nguồn được tổ chức khoa học với các module thực hiện các nhiệm vụ chuyên biệt:")
    add_bullets(doc, [
        "config.py: Định nghĩa các tham số cấu hình chung như kích thước ảnh đầu vào (64x128), tham số HOG (9 hướng, ô 8x8, khối 2x2), đường dẫn các thư mục và mô hình lưu trữ.",
        "features.py: Cung cấp hàm tiền xử lý ảnh (chuyển grayscale, cân bằng lược đồ xám bằng cv2.equalizeHist) và trích xuất vector đặc trưng HOG từ ảnh thô.",
        "detector.py: Định nghĩa lớp HOGSVMPedestrianDetector sử dụng mô hình HOG + SVM có sẵn của OpenCV. Đồng thời chứa hàm hậu xử lý NMS và hàm vẽ kết quả vẽ bounding box, nhãn độ tin cậy và FPS.",
        "custom_detector.py: Định nghĩa lớp CustomHOGSVMSlidingWindowDetector sử dụng kỹ thuật quét cửa sổ trượt (Sliding Window) đa tỷ lệ kết hợp với mô hình SVM tự huấn luyện (custom_hog_svm.pkl).",
        "download_inria_hf_subset.py: Tự động tải subset của bộ dữ liệu INRIA Person Dataset trên Hugging Face để phục vụ huấn luyện, lưu các ảnh dương (positive) và cắt ngẫu nhiên các ảnh âm (negative).",
        "train_custom_svm.py: Đọc tập dữ liệu mẫu, trích xuất đặc trưng HOG, phân chia tập train/test theo tỷ lệ 80/20 phân tầng, chuẩn hóa dữ liệu bằng StandardScaler và huấn luyện mô hình Linear SVM qua thư viện Scikit-Learn.",
        "evaluate_custom_svm.py: Đánh giá hiệu năng mô hình trên toàn bộ tập dữ liệu.",
        "visualize_gradient_hog.py: Đọc một ảnh kiểm thử, tính toán và xuất các ảnh trung gian Sobel Magnitude, Orientation và HOG visualization.",
        "run_realtime.py: Chương trình điều khiển chính kết nối Webcam hoặc đọc Video, gọi các lớp detector để dự đoán và hiển thị giao diện đồ họa."
    ])

    doc.add_heading("3.2 Luồng xử lý dữ liệu và Thuật toán dò tìm đa tỷ lệ (Multi-scale Sliding Window)", level=2)
    doc.add_paragraph(
        "Đối với mô hình tự huấn luyện (Custom SVM), vì nó chỉ nhận đầu vào cố định là ảnh 64x128, "
        "nhóm đã thiết lập giải thuật quét ảnh đa tỷ lệ để phát hiện người đi bộ với kích thước bất kỳ trong khung hình:"
    )
    add_numbered(doc, [
        "Tiền xử lý: Nhận frame từ video/webcam, resize ảnh về chiều rộng tối đa (ví dụ 360 pixel) để tăng tốc độ quét và giảm gánh nặng tính toán cho CPU.",
        "Khởi tạo danh sách các tỷ lệ ảnh (Scales): Tạo các hệ số thu phóng ảnh (ví dụ: 1.0, 0.75, 0.55, 0.40). Với mỗi hệ số, ảnh được resize tương ứng.",
        "Quét cửa sổ trượt (Sliding Window): Trên mỗi ảnh ở mỗi scale, sử dụng một cửa sổ quét kích thước cố định 64x128. Cửa sổ này dịch chuyển theo chiều ngang và dọc với bước nhảy window_step (ví dụ: 48 pixel).",
        "Trích xuất đặc trưng HOG cho cửa sổ và chuyển qua Pipeline học máy (chuẩn hóa StandardScaler và dự đoán bằng LinearSVC). Lấy điểm phân lớp từ hàm quyết định decision_function. Nếu điểm này lớn hơn decision_threshold, lưu lại tọa độ cửa sổ cùng điểm tin cậy.",
        "Hậu xử lý Non-Maximum Suppression (NMS) với ngưỡng IoU = 0.35 để loại bỏ các box trùng lặp trên cùng một đối tượng. Tiếp theo, chạy giải thuật loại bỏ các box nhỏ nằm phần lớn bên trong box lớn hơn (tỷ lệ chứa > 70%).",
        "Hiển thị: Vẽ bounding box kèm nhãn độ tin cậy và FPS lên khung hình."
    ])

    # CHƯƠNG 4
    doc.add_heading("CHƯƠNG 4: THÍ NGHIỆM VÀ HUẤN LUYỆN MÔ HÌNH CUSTOM", level=1)
    
    doc.add_heading("4.1 Thu thập dữ liệu ngoài từ INRIA Person Dataset", level=2)
    doc.add_paragraph(
        "Theo tiêu chí mở rộng và quy định sử dụng dữ liệu ngoài, dự án không lấy ảnh từ video demo để huấn luyện "
        "để tránh hiện tượng quá khớp (overfitting). Thay vào đó, nhóm đã kết nối trực tiếp với mirror của "
        "INRIA Person Dataset trên Hugging Face bằng kịch bản download_inria_hf_subset.py:"
    )
    add_bullets(doc, [
        "Ảnh Dương (Positive): Tải 300 ảnh chứa người đi bộ, tự động cắt và đưa về kích thước chuẩn 64x128 pixel.",
        "Ảnh Âm (Negative): Tải 300 ảnh phong cảnh không chứa người đi bộ. Kịch bản tự động cắt ngẫu nhiên (random crop) 2 vùng ảnh kích thước 64x128 trên mỗi bức ảnh nền để làm mẫu âm.",
        "Kết quả thu được tập dữ liệu huấn luyện cân bằng gồm 600 mẫu (300 Positive và 300 Negative)."
    ])
    add_dataset_table(doc)

    doc.add_heading("4.2 Cấu hình huấn luyện", level=2)
    doc.add_paragraph("Nhóm sử dụng thư viện scikit-learn để xây dựng luồng huấn luyện (Pipeline) tự động hóa:")
    add_bullets(doc, [
        "Chuẩn hóa dữ liệu: Sử dụng StandardScaler để đưa các giá trị đặc trưng HOG về phân phối có trung bình bằng 0 và độ lệch chuẩn bằng 1 giúp mô hình SVM hội tụ nhanh và ổn định.",
        "Mô hình: Bộ phân loại LinearSVC với hằng số phạt sai số C=1.0, trọng số lớp tự động cân bằng (class_weight='balanced'), và số vòng lặp tối đa là 20,000 để đảm bảo hội tụ.",
        "Phân chia dữ liệu: Chia ngẫu nhiên 80% cho tập huấn luyện (480 mẫu) và 20% cho tập kiểm thử (120 mẫu), áp dụng Stratified Split để cân bằng tỉ lệ nhãn."
    ])

    # CHƯƠNG 5
    doc.add_heading("CHƯƠNG 5: KẾT QUẢ THỰC NGHIỆM VÀ THẢO LUẬN", level=1)
    
    doc.add_heading("5.1 Các kết quả hình ảnh xử lý trung gian", level=2)
    doc.add_paragraph(
        "Để minh họa trực quan cho các bước xử lý ảnh trung gian trong giải thuật, các kết quả hình ảnh "
        "đã được trích xuất và hiển thị như dưới đây:"
    )
    add_figure(doc, "outputs/input_resized.jpg", "Hình 1. Ảnh đầu vào sau khi resize về kích thước chuẩn 64x128.", 2.2)
    add_figure(doc, "outputs/gradient_magnitude.jpg", "Hình 2. Biên độ Gradient Sobel (Gradient Magnitude).", 2.2)
    add_figure(doc, "outputs/gradient_orientation.jpg", "Hình 3. Hướng Gradient Sobel (Gradient Orientation).", 2.2)
    add_figure(doc, "outputs/hog_visualization.jpg", "Hình 4. Trực quan hóa đặc trưng HOG dưới dạng biểu đồ hướng trên từng cell.", 2.2)

    doc.add_heading("5.2 Đánh giá định lượng mô hình Custom SVM", level=2)
    doc.add_paragraph(
        "Mô hình huấn luyện Custom SVM đạt kết quả cực kỳ ấn tượng, chi tiết kết quả được tổng hợp trong bảng bên dưới:"
    )
    add_result_table(doc)

    doc.add_heading("5.3 So sánh mô hình OpenCV Pretrained và Custom Sliding-Window", level=2)
    doc.add_paragraph(
        "Nhóm tiến hành đối chiếu, so sánh ưu nhược điểm của hai phương pháp tiếp cận được tích hợp trong dự án:"
    )
    add_comparison_table(doc)

    # CHƯƠNG 6
    doc.add_heading("CHƯƠNG 6: NHẬN XÉT VÀ ĐÁNH GIÁ", level=1)
    
    doc.add_heading("6.1 Ưu điểm của phương pháp", level=2)
    add_bullets(doc, [
        "Tính khoa học: Phương pháp HOG + SVM thể hiện trực quan cách các thuật toán xử lý ảnh số sử dụng toán tử vi phân để chuyển đổi dữ liệu pixel thô sang dữ liệu đặc trưng hình dạng có ý nghĩa vật lý rõ ràng.",
        "Tính thực tiễn: Chương trình demo chính sử dụng OpenCV pretrained chạy mượt mà trên máy tính cá nhân thông thường mà không cần phần cứng GPU hỗ trợ.",
        "Độ liêm chính học thuật: Toàn bộ tập huấn luyện custom được tách biệt hoàn toàn khỏi video chạy demo thực tế. Nhóm tự tải và thiết lập bộ tham số thực nghiệm (kích thước bước nhảy window_step, ngưỡng ra quyết định decision_threshold, ngưỡng NMS)."
    ])

    doc.add_heading("6.2 Hạn chế và nguyên nhân", level=2)
    add_bullets(doc, [
        "Tốc độ của Custom model: Bộ phát hiện tự huấn luyện chạy khá chậm (~0.5 - 2 FPS). Nguyên nhân là do phép quét sliding window đa tỷ lệ được thực thi qua các vòng lặp lồng nhau trong Python, đồng thời việc trích xuất đặc trưng HOG cho hàng trăm cửa sổ quét trên mỗi frame tạo ra gánh nặng tính toán lớn. Để khắc phục, nhóm đã đưa ra giải pháp giảm kích thước tối đa của ảnh đầu vào và tăng bước nhảy cửa sổ quét.",
        "Nhạy cảm với bối cảnh: Cả mô hình OpenCV và mô hình Custom đều có thể phát sinh lỗi phát hiện nhầm (False Positive) khi nền ảnh chứa quá nhiều cạnh sắc thẳng đứng (như hàng rào, cột điện, thân cây) do chúng có đặc trưng phân bố HOG tương đối giống với cấu trúc dọc của cơ thể người."
    ])

    # CHƯƠNG 7
    doc.add_heading("CHƯƠNG 7: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)
    
    doc.add_heading("7.1 Kết luận", level=2)
    doc.add_paragraph(
        "Dự án cuối kỳ môn Nhập môn Xử lý ảnh số - Chủ đề 2 đã hoàn thành xuất sắc toàn bộ các mục tiêu đặt ra. "
        "Chúng tôi đã trình bày rõ ràng cơ sở lý thuyết toán học (Image Gradient, Sobel, HOG, SVM); "
        "xây dựng thành công ứng dụng demo hoạt động ổn định và mượt mà trên video/webcam; vẽ bounding box kèm nhãn realtime; "
        "đồng thời tích hợp thành công phần tải tập dữ liệu ngoài INRIA và huấn luyện mô hình Custom SVM đạt độ chính xác cao."
    )

    doc.add_heading("7.2 Hướng phát triển tương lai", level=2)
    add_bullets(doc, [
        "Tối ưu hóa mã nguồn: Sử dụng thư viện Numba hoặc Cython để biên dịch mã quét sliding window của Python sang mã máy, giúp tăng tốc độ thực thi của mô hình Custom.",
        "Khai phá mẫu âm khó (Hard Negative Mining): Thu thập các vùng ảnh nền bị nhận diện nhầm là người (False Positives), đưa chúng ngược trở lại tập dữ liệu huấn luyện để dạy mô hình SVM nhận diện tốt hơn, tránh lỗi nhiễu.",
        "Tích hợp mô hình học sâu: Nghiên cứu tích hợp thêm bộ phát hiện dựa trên mạng nơ-ron tích chập (YOLO) để tiến hành so sánh đối chiếu hiệu năng giữa xử lý ảnh truyền thống và học sâu hiện đại."
    ])

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(OUTPUT)
    print(f"Saved report: {OUTPUT}")


if __name__ == "__main__":
    build_report()
